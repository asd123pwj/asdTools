from BaseClasses.BaseModel import BaseModel
import docx
import os


class DocBase(BaseModel):
    """
    用于文档处理。

    Args:
        BaseModel (_type_): _description_
    """
    def __init__(self, docx_path:str, **kwargs) -> None:
        super().__init__(**kwargs)
        self.docx_paths = self.get_path_from_dir(docx_path)

    def convert_docx2pdf(self, 
                         docx_path:str, 
                         output_file:str="", 
                         output_dir:str="") -> str:
        if output_file == "":
            output_file = f"{self.convert_path2name(docx_path)}.pdf"
        if output_dir == "":
            output_dir = self.log_dir
        output_path = os.path.join(output_dir, output_file)
        docx_path = os.path.abspath(docx_path)
        output_path = os.path.abspath(output_path)
        # 方案1
        # from docx2pdf import convert
        # convert(docx_path, output_path)

        # 方案2
        # import win32com.client as win32
        # word = win32.DispatchEx("Word.Application")
        # doc = word.Documents.Open(docx_path)
        # doc.SaveAs(output_path, FileFormat=17)
        # doc.Close()
        # word.Quit()

        # 方案3
        # import subprocess
        # subprocess.call(['unoconv', '-f', 'pdf', docx_path])

        return output_path

    def read(self) -> None:
        for path in self.docx_paths:
            if self.check_ext(path, ["doc", "docx"]):
                self.read_docx(path)

    def read_docx(self, path:str) -> None:
        print(path)
        doc = docx.Document(path)
        content_para = []
        for ele in doc.paragraphs:
            txt = "".join(ele.text.split())
            if txt != "" and not ele.text in content_para:
                content_para.append(ele.text)
                print(ele.text)

        content_table = []
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    txt = "".join(cell.text.split())
                    if txt != "" and not cell.text in content_table:
                        content_table.append(cell.text)
                        print(cell.text)

        content_textbox = []
        children = doc.element.body.iter()
        for element in children:
            if element.tag.endswith('txbx'):
                text = ""
                for ele in element.iter():   
                    if ele.tag.endswith('main}r'):
                        text += f"{ele.text} "
                    
                text = " ".join(text.split())
                content_textbox.append(text)
                # print(text)

        dct = {f'{i+1}': f'{content_textbox[i]}' for i in range(len(content_textbox))}
        print(content_textbox)



        0


if __name__ == "__main__":
    # 打开文档
    docx_reader = DocBase("./data")
    docx_reader.read()

    
    # doc = docx.Document("./data/1.docx")

    # for para in doc.paragraphs:
    #     print(para.text)