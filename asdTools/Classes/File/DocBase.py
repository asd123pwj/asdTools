from asdTools.Classes.Base.BaseModel import BaseModel
import PyPDF2
import tempfile
import docx


class DocBase(BaseModel):
    def __init__(self, **kwargs) -> None:
        super().__init__(**kwargs)

    def load_docx(self, path:str):
        """Load a docx file.

        Args:
            path (str): The path to the docx file.

        Returns:
            Document: A docx.Document object.
        """
        return docx.Document(path)

    def read_para(self, doc) -> list:
        """Read paragraphs from a docx file.

        Args:
            doc (Document): A docx.Document object.

        Returns:
            list: A list of paragraphs.
        """
        content_para = []
        try:
            for ele in doc.paragraphs:
                txt = "".join(ele.text.split())
                if txt != "" and not ele.text in content_para:
                    content_para.append(ele.text)
        except:
            pass
        return content_para

    def read_table(self, doc) -> list:
        """Read tables from a docx file.

        Args:
            doc (Document): A docx.Document object.

        Returns:
            list: A list of tables.
        """
        content_table = []
        try:
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        txt = "".join(cell.text.split())
                        if txt != "" and not cell.text in content_table:
                            content_table.append(cell.text)
        except:
            pass
        return content_table

    def read_textbox(self, doc) -> list:
        """Read text boxes from a docx file.

        Args:
            doc (Document): A docx.Document object.

        Returns:
            list: A list of text boxes.
        """
        content_textbox = []
        try:
            children = doc.element.body.iter()
            for element in children:
                if element.tag.endswith('txbx'):
                    text = ""
                    for ele in element.iter():   
                        if ele.tag.endswith('main}r'):
                            text += f"{ele.text} "
                    text = " ".join(text.split())
                    content_textbox.append(text)
        except:
            pass
        return content_textbox

    def read_pdf(self, path:str):
        with open(path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            
            content = ""
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                c = page.extract_text()
                aspose_mark = r"Created with an evaluation copy of Aspose.Words. To discover the full versions \nof our APIs please visit: https://products.aspose.com/words/Evaluation Only. Created with Aspose.Words. Copyright 2003-2023 Aspose Pty \nLtd."
                if "Aspose.Words" in c:
                    res = c.split("\n")
                    res = res[3:]
                    c = "\n".join(res)
                c = c.replace("  ", "#shuangkongge")
                c = c.replace(" ", "")
                c = c.replace("#shuangkongge", " ")
                content += c
            return content

    def read_docx(self, path:str) -> str:
        doc = self.load_docx(path)
        content_para = self.read_para(doc)
        content_table = self.read_table(doc)
        content_textbox = self.read_textbox(doc)
        content = content_para + content_table + content_textbox
        content = "\n".join(content)
        if content == "":
            content = self.read_as_pdf(path)
        return content

    def read_as_pdf(self, path:str) -> str:
        ### comment below
        if not self.isAuthor():
            import os
            current_file_path = os.path.abspath(__file__)
            warning = "Considering doc2pdf needing aspose which is large (60M+), and most doc type is uncommon. "
            warning += "I dont add it to requirements.txt. "
            warning += "Please use `pip install aspose-words` to install aspose. "
            warning += f"And comments these code in {current_file_path}. "
            warning += "You can CTRL+F to find 'comment below' to find code."
            self.warning(warning)
            self.pause()
        ### comment above
        import aspose.words as aw
        with tempfile.NamedTemporaryFile(prefix='temp_', suffix='.pdf', delete=False) as temp_file:
            doc = aw.Document(path)
            temp_path = temp_file.name
        doc.save(temp_path)
        content = self.read_docxs(temp_path)
        self.remove(temp_path)
        return content

    def read_docxs(self, path:str) -> str:
        type = self.get_ext_of_file(path)
        if type == "docx":
            content = self.read_docx(path)
        elif type == "doc":
            content = self.read_as_pdf(path)
        elif type == "pdf":
            content = self.read_pdf(path)

        return content
