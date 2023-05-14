from asdTools.Classes.Base.BaseModel import BaseModel
import docx


class DocxBase(BaseModel):
    def __init__(self, **kwargs) -> None:
        super().__init__(**kwargs)

    @staticmethod
    def load_docx(path: str):
        """Load a docx file.

        Args:
            path (str): The path to the docx file.

        Returns:
            Document: A docx.Document object.
        """
        return docx.Document(path)

    @staticmethod
    def read_para(doc) -> list:
        """Read paragraphs from a docx file.

        Args:
            doc (Document): A docx.Document object.

        Returns:
            list: A list of paragraphs.
        """
        content_para = []
        for ele in doc.paragraphs:
            txt = "".join(ele.text.split())
            if txt != "" and not ele.text in content_para:
                content_para.append(ele.text)
        return content_para

    @staticmethod
    def read_table(doc) -> list:
        """Read tables from a docx file.

        Args:
            doc (Document): A docx.Document object.

        Returns:
            list: A list of tables.
        """
        content_table = []
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    txt = "".join(cell.text.split())
                    if txt != "" and not cell.text in content_table:
                        content_table.append(cell.text)
        return content_table

    @staticmethod
    def read_textbox(doc) -> list:
        """Read text boxes from a docx file.

        Args:
            doc (Document): A docx.Document object.

        Returns:
            list: A list of text boxes.
        """
        content_textbox = []
        children = doc.element.body.iter()
        for element in children:
            if element.tag.endswith('txbx'):
                text = ""
                for ele in element.iter():   
                    if ele.tag.endswith('main}r'):
                        text += f"{ele.text} "
                text = " ".join(text.split())
                content_textbox.append(text)
        return content_textbox

    def read_docx(self, path: str) -> list:
        """Read paragraphs, tables, and text boxes from a docx file.

        Args:
            path (str): The path to the docx file.

        Returns:
            list: A list of paragraphs, tables, and text boxes.
        """
        doc = self.load_docx(path)
        content_para = self.read_para(doc)
        content_table = self.read_table(doc)
        content_textbox = self.read_textbox(doc)
        content = content_para + content_table + content_textbox
        return content
